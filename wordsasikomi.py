import PySimpleGUI as sg
from docx import Document
from tkinter import Tk
from tkinter.filedialog import askopenfilename, asksaveasfilename
from datetime import datetime
import re
import sys
import os # Add this import

# --- Debug Log File ---
DEBUG_LOG_FILE = "debug_log.txt"
def log_debug(message):
    with open(DEBUG_LOG_FILE, "a", encoding="utf-8") as f:
        f.write(message + "\n")

# --- Wordファイル選択 ---
def select_word_file():
    Tk().withdraw()
    file_path = askopenfilename(title="Wordテンプレートを選択", filetypes=[("Word files", "*.docx")])
    return file_path

# --- 段落内のrunを結合してテキストを取得 ---
def get_paragraph_text(paragraph):
    """段落内のすべてのrunを結合してテキストを取得"""
    return ''.join(run.text for run in paragraph.runs)

# --- 段落内のプレースホルダーを適切に置換 ---
def replace_in_paragraph(paragraph, placeholder, replacement):
    """段落内のプレースホルダーを置換（runの境界を考慮）"""
    full_text = get_paragraph_text(paragraph)
    
    # プレースホルダーが存在するか確認
    if "{{" + placeholder + "}}" not in full_text: # Corrected f-string
        return False
    
    # 新しいテキストを作成
    new_text = full_text.replace("{{" + placeholder + "}}", replacement)
    
    # 既存のrunのスタイルを保持しつつ置換
    if paragraph.runs:
        # 最初のrunのスタイルを保持
        first_run = paragraph.runs[0]
        # すべてのrunをクリア
        for run in paragraph.runs:
            run.text = ""
        # 最初のrunに新しいテキストを設定
        first_run.text = new_text
    else:
        # runがない場合は新しく作成
        paragraph.text = new_text
    
    return True

# --- プレースホルダー検出関数（改善版） ---
def find_placeholders(doc):
    placeholders = set()  # 重複を避けるためにセットを使用
    placeholder_instances = []
    
    # 段落内のプレースホルダー検索
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        log_debug(f"Paragraph {i} raw text (repr): {repr(text)}") # Debug print to file
        # プレースホルダーのパターン {{名前}} を検索
        matches = re.findall(r'{{(.*?)}}', text)
        for ph_name in matches:
            placeholders.add(ph_name.strip())
            placeholder_instances.append({
                "para_index": i, 
                "placeholder_name": ph_name.strip(),
                "type": "paragraph"
            })
    
    # テーブル内のプレースホルダー検索
    for i, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            log_debug(f"Table[{i}] Row[{row_idx}] has {len(row.cells)} cells.") # Debug print to file
            for cell_idx, cell in enumerate(row.cells):
                log_debug(f"Table[{i}] Row[{row_idx}] Cell[{cell_idx}] has {len(cell.paragraphs)} paragraphs.") # Debug print to file
                for para_idx, para in enumerate(cell.paragraphs):
                    text = para.text
                    log_debug(f"Table[{i}] Row[{row_idx}] Cell[{cell_idx}] Para[{para_idx}] raw text (repr): {repr(text)}") # Debug print to file
                    matches = re.findall(r'{{(.*?)}}', text)
                    for ph_name in matches:
                        placeholders.add(ph_name.strip())
                        placeholder_instances.append({
                            "table_index": i,
                            "row_idx": row_idx,
                            "cell_idx": cell_idx,
                            "para_idx": para_idx,
                            "placeholder_name": ph_name.strip(),
                            "type": "table"
                        })
    
    return list(placeholders), placeholder_instances

# --- 置換実行関数 ---
def replace_placeholders(doc, replacements, placeholder_instances):
    # 段落内の置換
    for instance in placeholder_instances:
        ph_name = instance["placeholder_name"]
        if ph_name in replacements:
            replacement_value = replacements[ph_name]
            
            if instance["type"] == "paragraph":
                para = doc.paragraphs[instance["para_index"]]
                para.text = para.text.replace(f"{{{{{ph_name}}}}}", replacement_value)
            
            elif instance["type"] == "table":
                log_debug(f"Processing table instance: {instance}") # Debug print to file
                table = doc.tables[instance["table_index"]]
                log_debug(f"Table rows: {len(table.rows)}, instance row_idx: {instance['row_idx']}") # Debug print to file
                row = table.rows[instance["row_idx"]]
                log_debug(f"Row cells: {row.cells}, instance cell_idx: {instance['cell_idx']}") # Debug print to file
                cell = row.cells[instance["cell_idx"]]
                para = cell.paragraphs[instance["para_idx"]]
                para.text = para.text.replace(f"{{{{{ph_name}}}}}", replacement_value)
    
    return doc

# メイン処理
def main():
    file_path = select_word_file()
    if not file_path:
        sg.popup("ファイルが選択されませんでした。")
        return
    
    # Wordファイルを開く
    try:
        doc = Document(file_path)
    except Exception as e:
        sg.popup(f"ファイルを開けませんでした: {e}")
        return
    
    # プレースホルダーの検出
    unique_placeholders, placeholder_instances = find_placeholders(doc)
    
    if not unique_placeholders:
        sg.popup("プレースホルダーが見つかりませんでした。nWord文書に {{プレースホルダー名}} の形式で記述してください。")
        return
    
    # GUIレイアウト作成
    layout = []
    
    # 入力フィールドの作成（重複なし）
    input_fields = []
    for ph_name in sorted(unique_placeholders):  # ソートして表示
        input_key = f"-PH-{ph_name}-"
        input_fields.append([sg.Text(f"{{{{ {ph_name} }}}}", size=(30, 1))])
        input_fields.append([sg.Input(key=input_key, size=(50, 1))])
    
    layout.append([sg.Column(input_fields, scrollable=True, vertical_scroll_only=True, size=(600, 400))])
    layout.append([sg.Button("差し替え"), sg.Button("終了")])
    
    window = sg.Window("Word差し込みツール", layout, finalize=True, size=(650, 550))
    
    # GUIイベントループ
    while True:
        event, values = window.read()
        
        if event in (sg.WINDOW_CLOSED, "終了"):
            break
            
        if event == "差し替え":
            # 未入力チェック
            empty_fields = []
            for ph_name in unique_placeholders:
                input_key = f"-PH-{ph_name}-"
                if not values[input_key].strip():
                    empty_fields.append(ph_name)
            
            if empty_fields:
                sg.popup(f"未入力の項目があります:n{', '.join(['{{'+ph+'}}' for ph in empty_fields])}")
                continue
            
            # 置換用の辞書を作成
            replacements = {}
            for ph_name in unique_placeholders:
                input_key = f"-PH-{ph_name}-"
                replacements[ph_name] = values[input_key]
            
            # 置換実行
            doc = replace_placeholders(doc, replacements, placeholder_instances)
            
            # 保存先を選択
            default_filename = f"output_{{datetime.now().strftime('%Y%m%d_%H%M%S')}}.docx"
            Tk().withdraw()
            save_path = asksaveasfilename(
                title="保存先を選択",
                initialfile=default_filename,
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")]
            )
            
            if save_path:
                try:
                    doc.save(save_path)
                    sg.popup(f"保存しました: {save_path}")
                except Exception as e:
                    sg.popup(f"保存に失敗しました: {e}")
    
    window.close()

if __name__ == "__main__":
    main()